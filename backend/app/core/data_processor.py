"""
Data Processor for FastMCP Server.

This module provides functionality for:
- File content extraction
- Data validation
- Batch processing
- Data transformation
"""

import os
import json
import csv
import logging
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
import pandas as pd
from fastapi import UploadFile
import PyPDF2
import docx
import openpyxl

logger = logging.getLogger(__name__)

class DataProcessor:
    """Handles data processing operations for files."""
    
    SUPPORTED_TYPES = {
        'txt': 'text',
        'json': 'json',
        'csv': 'csv',
        'xlsx': 'excel',
        'xls': 'excel',
        'pdf': 'pdf',
        'docx': 'word',
        'doc': 'word'
    }
    
    def __init__(self):
        """Initialize the data processor."""
        self.extractors = {
            'text': self._extract_text,
            'json': self._extract_json,
            'csv': self._extract_csv,
            'excel': self._extract_excel,
            'pdf': self._extract_pdf,
            'word': self._extract_word
        }
        
    async def extract_content(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Extract content from a file based on its type.
        
        Args:
            file_path: Path to the file
            file_type: Type of the file (extension)
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        try:
            file_type = file_type.lower()
            if file_type not in self.SUPPORTED_TYPES:
                raise ValueError(f"Unsupported file type: {file_type}")
                
            processor_type = self.SUPPORTED_TYPES[file_type]
            extractor = self.extractors[processor_type]
            
            content = await extractor(file_path)
            return {
                "content": content,
                "type": processor_type,
                "extracted_at": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {str(e)}")
            raise
            
    async def _extract_text(self, file_path: str) -> str:
        """Extract content from a text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
            
    async def _extract_json(self, file_path: str) -> Dict[str, Any]:
        """Extract content from a JSON file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
            
    async def _extract_csv(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract content from a CSV file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            return list(reader)
            
    async def _extract_excel(self, file_path: str) -> Dict[str, Any]:
        """Extract content from an Excel file."""
        df = pd.read_excel(file_path, sheet_name=None)
        return {
            sheet: df[sheet].to_dict('records')
            for sheet in df.keys()
        }
        
    async def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract content from a PDF file."""
        with open(file_path, 'rb') as f:
            pdf = PyPDF2.PdfReader(f)
            return {
                "pages": len(pdf.pages),
                "content": [
                    page.extract_text()
                    for page in pdf.pages
                ]
            }
            
    async def _extract_word(self, file_path: str) -> Dict[str, Any]:
        """Extract content from a Word document."""
        doc = docx.Document(file_path)
        return {
            "paragraphs": [p.text for p in doc.paragraphs],
            "tables": [
                [
                    [cell.text for cell in row.cells]
                    for row in table.rows
                ]
                for table in doc.tables
            ]
        }
        
    async def validate_data(self, data: Any, schema: Dict[str, Any]) -> Dict[str, Any]:
        """
        Validate data against a schema.
        
        Args:
            data: Data to validate
            schema: Validation schema
            
        Returns:
            Dictionary containing validation results
        """
        try:
            # Basic type validation
            if not isinstance(data, type(schema.get('type', type(data)))):
                return {
                    "valid": False,
                    "errors": [f"Expected type {schema.get('type')}, got {type(data)}"]
                }
                
            # Additional validation based on schema
            errors = []
            
            # Required fields
            if 'required' in schema:
                for field in schema['required']:
                    if field not in data:
                        errors.append(f"Missing required field: {field}")
                        
            # Field type validation
            if 'properties' in schema:
                for field, field_schema in schema['properties'].items():
                    if field in data:
                        field_type = field_schema.get('type')
                        if field_type and not isinstance(data[field], eval(field_type)):
                            errors.append(
                                f"Field {field} should be of type {field_type}"
                            )
                            
            return {
                "valid": len(errors) == 0,
                "errors": errors
            }
            
        except Exception as e:
            logger.error(f"Error validating data: {str(e)}")
            raise
            
    async def process_batch(
        self,
        files: List[UploadFile],
        processor_type: str,
        schema: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Process a batch of files.
        
        Args:
            files: List of files to process
            processor_type: Type of processing to apply
            schema: Optional validation schema
            
        Returns:
            Dictionary containing processing results
        """
        results = []
        errors = []
        
        for file in files:
            try:
                # Save file temporarily
                temp_path = f"data/tmp/{file.filename}"
                with open(temp_path, "wb") as f:
                    content = await file.read()
                    f.write(content)
                    
                # Extract content
                file_type = file.filename.rsplit('.', 1)[1].lower()
                content = await self.extract_content(temp_path, file_type)
                
                # Validate if schema provided
                if schema:
                    validation = await self.validate_data(content, schema)
                    if not validation['valid']:
                        errors.append({
                            "file": file.filename,
                            "errors": validation['errors']
                        })
                        continue
                        
                results.append({
                    "file": file.filename,
                    "content": content
                })
                
            except Exception as e:
                errors.append({
                    "file": file.filename,
                    "error": str(e)
                })
                
            finally:
                # Cleanup
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                    
        return {
            "processed": len(results),
            "failed": len(errors),
            "results": results,
            "errors": errors
        }

# Singleton instance
data_processor = DataProcessor() 